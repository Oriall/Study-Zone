from flask import (
    Flask, 
    render_template, 
    send_file, 
    request, 
    jsonify, 
    redirect, url_for, 
    Response,
)
from pathlib import Path
from flask import session, send_from_directory
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import threading
import requests
import os
import markdown
from docx import Document
import openai
from docx.shared import Pt
from twilio.rest import Client
import keyboard
import mysql.connector
from flask_cors import CORS
from flask_socketio import SocketIO
from functools import wraps
import datetime
from datetime import datetime
from dotenv import load_dotenv
import re

load_dotenv()
app = Flask(__name__)
BASE_DIR = Path(__file__).resolve().parent
DATA_FILE_PATH2 = BASE_DIR / 'Capture_Folder(StudyZone)'
# # Define the website to block (e.g., Facebook)
# website_to_block = "www.facebook.com"

# # Define the path to the hosts file on Windows
# hosts_path = r"C:\Windows\System32\drivers\etc\hosts"

# # Add an entry to the hosts file to block the website
# def block_website(website):
#     with open(hosts_path, "a") as hosts_file:
#         hosts_file.write(f"127.0.0.1 {website}\n")
#     return f"{website} is blocked."

# # Unblock the website (remove the entry from the hosts file)
# def unblock_website(website):
#     with open(hosts_path, "r+") as hosts_file:
#         lines = hosts_file.readlines()
#         hosts_file.seek(0)
#         for line in lines:
#             if not website in line:
#                 hosts_file.write(line)
#         hosts_file.truncate()
#     return f"{website} is unblocked."
# Đường dẫn đến thư mục cần tạo
directory = "data_get"
file_path_data = os.path.join(directory, "data.txt")

# Kiểm tra nếu thư mục chưa tồn tại, thì tạo mới
if not os.path.exists(directory):
    os.makedirs(directory)
    print(f'Thư mục "{directory}" đã được tạo.')
else:
    print(f'Thư mục "{directory}" đã tồn tại.')

# Kiểm tra và tạo file nếu chưa tồn tại
if not os.path.exists(file_path_data):
    with open(file_path_data, 'w') as file:
        file.write('')  # Tạo file rỗng
    print(f'File "{file_path_data}" đã được tạo.')
else:
    print(f'File "{file_path_data}" đã tồn tại.')
    
@app.route("/")
@app.route("/home")
def home():
    return render_template("index.html")

@app.route("/library")
def library():
    return render_template("library.html")

@app.route("/info_page")
def info_page():
    return render_template("info_page.html")


@app.route("/landing")
def landing():
    return render_template("landing.html")
@app.route("/setting2-result")
def setting2_result():
    # if 'html_content' in session:
    #     html_content = session['html_content']
    return render_template("setting2-result.html")
@app.route("/wait")
def wait():
    return render_template("wait.html")



UPLOAD_FOLDER3 = 'data_get'
app.config['UPLOAD_FOLDER3'] = UPLOAD_FOLDER3

@app.route("/library_parent")
def library_parent():
    learner_id_2 = session.get('acc', '')
    cursor.execute("SELECT * FROM screenshots WHERE LearnerId = %s", (learner_id_2,))
    filenames = cursor.fetchall()
    data_file_path = os.path.join(app.config['UPLOAD_FOLDER3'], 'data.txt')
    with open(data_file_path, 'r') as file:
        data = file.read()
    return render_template("library_parent.html", filenames=filenames, data=data)

@app.route("/setting4")
def setting4():
    return render_template("setting4.html")
@app.route("/setting4_result")
def setting4_result():
    return render_template("setting4_result.html")
@app.route("/macros")
def macros():
    return render_template("macros.html")

@app.route("/setting2")
def setting2():
    return render_template("setting2.html")
@app.route("/add_student")
def add_student():
    return render_template("add_student.html")
@app.route("/page-study")
def page_study():
    return render_template("page-study.html")

    ##return render_template("main.html",result=result)
@app.route("/static-page")
def static_page():
    return render_template("static-page.html")
@app.route("/setting3")
def setting3():
    return render_template("setting3.html")

@app.route("/add_teacher")
def add_teacher():
    return render_template("add_teacher.html")



@app.route("/logup")
def logup():
    return render_template("logup.html")

@app.route("/download")
def download():
    return render_template("download.html")


####### DATABSE ############
mydb = mysql.connector.connect(
    host=os.environ.get('DATABASE_HOST'),
    user=os.environ.get('DATABASE_USER'),
    password=os.environ.get('DATABASE_PASS'),
    database=os.environ.get('DATABASE'),
)

# Tạo cursor
mycursor = mydb.cursor()

# code_results Table
mycursor.execute("SHOW TABLES LIKE 'code_results'")
result = mycursor.fetchone()
if not result:
    mycursor.execute("""
        CREATE TABLE code_results (
            ID INT AUTO_INCREMENT PRIMARY KEY,
            LearnerID INT,
            Type INT,
            Message TEXT,
            Code TEXT,
            Date DATE,
            Time TIME,
            NameCode TEXT
        )
    """)
    print("Table created successfully.")
else:
    print("Table already exists.")

# Classes Table
mycursor.execute("SHOW TABLES LIKE 'Classes'")
result = mycursor.fetchone()
if not result:
    mycursor.execute("""
        CREATE TABLE Classes (
            ClassID INT AUTO_INCREMENT PRIMARY KEY,
            ClassName VARCHAR(10),
            TeacherID INT,
            Note VARCHAR(255)
        )
    """)
    print("Table 'Classes' created successfully.")
else:
    print("Table 'Classes' already exists.")

# Lấy số lượng bản ghi hiện tại trong bảng "Classes"
mycursor.execute("SELECT COUNT(*) FROM Classes")
count = mycursor.fetchone()[0]
#------------------------ DATA----------------------
if count < 6:
    classNames = ["10A1", "10A2", "10A3", "11A1", "11A2", "11A3"]
    teacherIDs = [1, 1, 1, 2, 2, 1]
    notes = ["", "", "", "", "", ""]

    for i in range(len(classNames)):
        sql = "INSERT INTO Classes (ClassName, TeacherID, Note) VALUES (%s, %s, %s)"
        val = (classNames[i], teacherIDs[i], notes[i])
        mycursor.execute(sql, val)

    mydb.commit()
    print("Data inserted into 'Classes' table successfully.")

# Teachers Table
mycursor.execute("SHOW TABLES LIKE 'Teachers'")
result = mycursor.fetchone()
if not result:
    mycursor.execute("""
        CREATE TABLE Teachers (
            TeacherID INT AUTO_INCREMENT PRIMARY KEY,
            Fullname VARCHAR(255),
            Email VARCHAR(255),
            Password VARCHAR(255)
        )
    """)
    print("Table 'Teachers' created successfully.")
else:
    print("Table 'Teachers' already exists.")
@app.route('/submit3', methods=['POST'])
def submit3():
    if request.method == 'POST':
        fullname = request.form['username']
        email = request.form['email']
        password = request.form['password']

        # Lưu thông tin vào cơ sở dữ liệu
        sql = "INSERT INTO Teachers (Fullname, Email, Password) VALUES (%s, %s, %s)"
        val = (fullname, email, password)
        mycursor.execute(sql, val)
        mydb.commit()

        return 'Đã lưu thông tin thành công!' 

@app.route('/submit4', methods=['POST'])
def submit4():
    if request.method == 'POST':
        class_name = request.form['class_name']
        fullname = request.form['fullname']
        note = request.form['note']

        # Tìm TeacherID tương ứng với Fullname
        cursor = mydb.cursor()
        cursor.execute("SELECT TeacherID FROM teachers WHERE Fullname = %s", (fullname,))
        teacher_id = cursor.fetchone()

        if teacher_id:
            teacher_id = teacher_id[0]

            # Chèn dữ liệu vào bảng Classes
            cursor.execute("INSERT INTO classes (ClassName, TeacherID, Note) VALUES (%s, %s, %s)",
                           (class_name, teacher_id, note))
            mydb.commit()
            cursor.close()
            return 'Đăng ký thành công!'
        else:
            return 'Không tìm thấy giáo viên có tên này!'


mycursor.execute("CREATE TABLE IF NOT EXISTS Learners (LearnerID INT AUTO_INCREMENT PRIMARY KEY, ClassID INT, Fullname VARCHAR(255), Email VARCHAR(255), Password VARCHAR(255), ParentName VARCHAR(255), ParentEmail VARCHAR(255), ParentPsw VARCHAR(255), Status INT)")
@app.route('/submit2', methods=['POST'])
def submit2():
    if request.method == 'POST':
        class_name = request.form['class_name']
        fullname = request.form['fullname']
        email = request.form['email']
        password = request.form['password']
        parent_name = request.form['parent_name']
        parent_email = request.form['parent_email']
        parent_psw = request.form['parent_password']

        # Xác định ClassID
        if '10A1' in class_name:
            class_id = 1
        elif '10A2' in class_name:
            class_id = 2
        elif '10A3' in class_name:
            class_id = 3
        elif '11A1' in class_name:
            class_id = 4
        elif '11A2' in class_name:
            class_id = 5
        elif '11A3' in class_name:
            class_id = 6
        else:
            class_id = 0  # Nếu không tìm thấy số lớp, gán 0

        # Lưu thông tin vào cơ sở dữ liệu
        sql = "INSERT INTO Learners (ClassID, Fullname, Email, Password, ParentName, ParentEmail, ParentPsw, Status) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)"
        val = (class_id, fullname, email, password, parent_name, parent_email, parent_psw, 0)
        mycursor.execute(sql, val)
        mydb.commit()

        return 'Đã lưu thông tin thành công!'
    
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Kiểm tra xem người dùng đã đăng nhập hay chưa
        if 'email2' not in session and 'email' not in session and 'email3' not in session:
            # Nếu không, điều hướng về trang login
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function
@app.route("/login", methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']

        # HỌC SINH
        sql = "SELECT * FROM Learners WHERE Email = %s AND Password = %s"
        val = (email, password)
        mycursor.execute(sql, val)
        account = mycursor.fetchone()
        global learner_id_2
        if account:
            session['email2'] = email
            session['fullname'] = account[2]
            session['acc'] = account[0]
            return redirect(url_for('join'))
        #PHỤ HUYNH
        sql_parent = "SELECT * FROM Learners WHERE ParentEmail = %s AND ParentPsw = %s"
        val_parent = (email, password)
        mycursor.execute(sql_parent, val_parent)
        parent_account = mycursor.fetchone()

        if parent_account:
            session['email3'] = email
            session['fullname'] = parent_account[5]
            session['acc'] = parent_account[0]
            return redirect(url_for('gs_parent'))
        #THẦY CO GIÁO
        sql_teacher = "SELECT * FROM Teachers WHERE Email = %s AND Password = %s"
        val_teacher = (email, password)
        mycursor.execute(sql_teacher, val_teacher)
        teacher_account = mycursor.fetchone()

        if teacher_account:
            session['email'] = email
            session['fullname'] = teacher_account[1]
            session['acc'] = teacher_account[0]
            return redirect(url_for('teacher_classes'))
        
        return 'Email hoặc mật khẩu không chính xác!'
    
    return render_template("login.html")


# Route để lấy email của học sinh tương ứng với giáo viên đang đăng nhập
@app.route('/get_teacher_emails', methods=['GET'])
def get_teacher_emails():
    # Lấy thông tin về email của giáo viên đang đăng nhập từ session
    teacher_email = session.get('email', None)
    if not teacher_email:
        return 'Vui lòng đăng nhập để truy cập.'

    # Lấy TeacherID của giáo viên đang đăng nhập
    sql_teacher_id = "SELECT TeacherID FROM Teachers WHERE Email = %s"
    mycursor.execute(sql_teacher_id, (teacher_email,))
    teacher_id_result = mycursor.fetchone()
    if not teacher_id_result:
        return 'Không tìm thấy thông tin về giáo viên.'

    teacher_id = teacher_id_result[0]

    # Lấy các ClassID mà giáo viên đó kiểm soát từ bảng Classes
    sql_teacher_classes = "SELECT ClassID FROM Classes WHERE TeacherID = %s"
    mycursor.execute(sql_teacher_classes, (teacher_id,))
    teacher_class_ids = [result[0] for result in mycursor.fetchall()]

    # Lấy các Email tương ứng từ bảng Learners
    teacher_emails = []
    # Lấy các Email và Fullname tương ứng từ bảng Learners
    for class_id in teacher_class_ids:
        sql_class_learners = "SELECT Email, Fullname FROM Learners WHERE ClassID = %s"
        mycursor.execute(sql_class_learners, (class_id,))
        learner_info = mycursor.fetchall()
        teacher_emails.extend(learner_info)

    # Tách riêng các giá trị Email và Fullname
    emails = [info[0] for info in teacher_emails]
    fullnames = [info[1] for info in teacher_emails]

    return render_template('teacher_emails.html', emails=emails, fullnames=fullnames)


@app.route('/teacher_classes')
@login_required
def teacher_classes():
    # Lấy TeacherID của giáo viên đang đăng nhập từ session
    teacher_email = session.get('email', None)
    if not teacher_email:
        return 'Vui lòng đăng nhập để truy cập.'

    # Lấy TeacherID của giáo viên đang đăng nhập
    sql_teacher_id = "SELECT TeacherID,Fullname FROM Teachers WHERE Email = %s"
    mycursor.execute(sql_teacher_id, (teacher_email,))
    teacher_id_result = mycursor.fetchone()
    if not teacher_id_result:
        return 'Không tìm thấy thông tin về giáo viên.'

    teacher_id = teacher_id_result[0]
    teacher_name = teacher_id_result[1]

    

    # Lấy các ClassName mà giáo viên đó phụ trách từ bảng Classes
    sql_teacher_classes = "SELECT ClassID,ClassName FROM Classes WHERE TeacherID = %s"
    mycursor.execute(sql_teacher_classes, (teacher_id,))
    teacher_classes = mycursor.fetchall()


    return render_template('teacher_classes.html', teacher_classes=teacher_classes, teacher_name=teacher_name)



def generate_report(email):
    sql = "SELECT Fullname, LearnerID FROM Learners WHERE Email = %s"
    val = (email,)
    mycursor.execute(sql, val)
    learner_info = mycursor.fetchone()

    if learner_info:
        fullname = learner_info[0]
        class_id = learner_info[1]

        sql_messages = "SELECT Message, Type, Code, Date, Time FROM code_results WHERE LearnerID = %s ORDER BY ID"
        val_messages = (class_id,)
        mycursor.execute(sql_messages, val_messages)
        messages = mycursor.fetchall()

        document = Document()
        document.add_heading('Báo cáo tổng hợp', 0).bold = True

        # Thêm phần "Họ và tên" theo định dạng mong muốn
        paragraph = document.add_paragraph()
        run = paragraph.add_run('Họ và tên: ')
        run.bold = True
        run.font.size = Pt(14)  # Sử dụng cỡ chữ 14
        run = paragraph.add_run(fullname)
        run.font.size = Pt(14)

        # Thêm phần "Ngày" theo định dạng mong muốn
        paragraph = document.add_paragraph()
        run = paragraph.add_run('Ngày: ')
        run.bold = True
        run.font.size = Pt(14)
        run = paragraph.add_run(datetime.now().strftime('%d/%m/%Y'))
        run.font.size = Pt(14)

        start_time = None  # Khởi tạo biến thời gian bắt đầu
        end_time = None    # Khởi tạo biến thời gian kết thúc
        dem=0
        for message in messages:
            if message[1] == 0:
                dem=dem+1
                document.add_heading('Bài ' + str(dem), level=1).bold = True
                document.add_paragraph(message[0], style='BodyText')  # Sử dụng kiểu chữ mặc định cho văn bản
                document.add_heading('Chương trình hoàn chỉnh:').bold = True
                document.add_paragraph(message[2], style='BodyText')  # Thêm phần Code
                
            elif message[1] == 1:
                document.add_heading('Các lỗi:', level=2).bold = True
                document.add_paragraph('[Error] ' + message[0], style='BodyText')  # Thêm chữ [Error] trước tin nhắn lỗi

            if message[1] == 1 and start_time is None:  # Nếu gặp message[1] == 1 đầu tiên của mỗi bài và chưa có thời gian bắt đầu
                start_time = message[4]  # Gán thời gian vào thời gian bắt đầu
            elif message[1] == 0:  # Nếu gặp message[1] == 0
                end_time = message[4]  # Gán thời gian vào thời gian kết thúc
                document.add_paragraph('Thời gian bắt đầu: ' + str(start_time), style='BodyText')  # Thêm thời gian bắt đầu vào báo cáo
                document.add_paragraph('Thời gian kết thúc: ' + str(end_time), style='BodyText')  # Thêm thời gian kết thúc vào báo cáo
                start_time = None  # Reset thời gian bắt đầu
                document.add_paragraph('----------------------------------------------------------------------------------------------------------------------')
        filename = 'Báo Cáo.docx'
        document.save(filename)
        return filename
    else:
        return None
    
@app.route('/average_completion_time/<learner_id>')
def average_completion_time(learner_id):
    try:
        # Truy vấn để lấy thông tin về thời gian hoàn thành các bài học
        sql_get_completion_time = """
            SELECT ID, Type, Time
            FROM code_results
            WHERE LearnerID = %s
            ORDER BY ID;
        """
        mycursor.execute(sql_get_completion_time, (learner_id,))
        results = mycursor.fetchall()

        # Xử lý dữ liệu để tính thời gian hoàn thành của mỗi bài học
        completion_times = []
        start_time = None
        count = 1
        for result in results:
            if result[1] == 1:  # Nếu là type 1 (lỗi)
                start_time = result[2]
            elif result[1] == 0 and start_time is not None:  # Nếu là type 0 (hoàn thành)
                end_time = result[2]
                time_difference = end_time - start_time
                completion_times.append((time_difference.total_seconds() // 60))  # Chuyển sang phút và lưu lại
                count += 1
                start_time = None

        
        average_completion_time = sum(completion_times) / len(completion_times) if len(completion_times) > 0 else 0

        return jsonify({'learner_id': learner_id, 'average_completion_time': average_completion_time})
    except Exception as e:
        return str(e)


@app.route("/gs_parent")
@login_required
def gs_parent():
    learner_id_2 = session.get('acc', '')
    mycursor.execute("SELECT * FROM Learners WHERE LearnerId = %s", (learner_id_2,))
    filename = mycursor.fetchone()
    sql_count_type_1 = """
        SELECT COUNT(*) AS Type1Count
        FROM code_results
        WHERE LearnerID = %s AND Type = 1 
    """
    mycursor.execute(sql_count_type_1, (learner_id_2,))
    learner_type_1_counts = mycursor.fetchone()[0]

    sql_count_type_0 = """
        SELECT COUNT(*) AS Type0Count
        FROM code_results
        WHERE LearnerID = %s AND Type = 0 
    """
    mycursor.execute(sql_count_type_0, (learner_id_2,))
    learner_type_0_counts = mycursor.fetchone()[0]
    response = requests.get(url_for('average_completion_time', learner_id=learner_id_2, _external=True))
    learner_average_times = response.json()['average_completion_time']
    return render_template("gs_parent.html", filename=filename, learner_type_1_counts=learner_type_1_counts, learner_type_0_counts=learner_type_0_counts, learner_average_times=learner_average_times)

@app.route("/admin_page")
def admin_page():
    mycursor.execute("SELECT * FROM Learners")
    learners = mycursor.fetchall()
    mycursor.execute("SELECT COUNT(learnerID) FROM learners")
    dem_hs = mycursor.fetchone()[0]
    mycursor.execute("SELECT COUNT(TeacherID) FROM teachers")
    dem_gv = mycursor.fetchone()[0]
    mycursor.execute("SELECT COUNT(ClassID) FROM classes")
    dem_lop = mycursor.fetchone()[0]
    return render_template("admin_page.html", learners=learners, dem_hs=dem_hs, dem_gv=dem_gv, dem_lop=dem_lop)


@app.route('/download_report')
def download_report():
    email = request.args.get('email')
    if email:
        # Tạo và trả về file báo cáo
        filename = generate_report(email)
        if filename:
            return send_file(filename, as_attachment=True)
        else:
            return 'Không thể tạo báo cáo.'
####### DATABSE ############
app.config['SECRET_KEY'] = 'khanhtoan123'
openai.api_key = os.environ.get('OPEN_AI_KEY')
socketio = SocketIO(app)
results = []

def get_result(class_id):
    code = session.get('code', '')
    
#     # Sử dụng API của ChatGPT để phân tích đề của code
    prompt = f"Đưa ra đề bài cho code(1 dòng)(bằng tiếng việt):\n\n{code}\n\n"

    response = openai.chat.completions.create(
        model="gpt-3.5-turbo", 
        messages=[
            {"role": "system", "content": prompt},
        ],
        max_tokens=1024
    )

    problem_statement = response.choices[0].message.content.strip()

    # Lưu kết quả, thời gian, và số lần click vào danh sách
    result = session.get('result', '')

    results.append({'code': code, 'problem_statement': problem_statement, 'result':result})

     # Gửi kết quả đến tất cả các clients đang kết nối
    socketio.emit('update_result', {'code': code, 'problem_statement': problem_statement, 'result': result})

    session['problem_statement'] = problem_statement

##------------------------------------------------------------------------------------------------------------##
    # Lấy thông tin về email của giáo viên đang đăng nhập từ session
    teacher_email = session.get('email', None)
    if not teacher_email:
        return 'Vui lòng đăng nhập để truy cập.'

    # Lấy TeacherID của giáo viên đang đăng nhập
    sql_teacher_id = "SELECT TeacherID,Fullname FROM Teachers WHERE Email = %s"
    mycursor.execute(sql_teacher_id, (teacher_email,))
    teacher_id_result = mycursor.fetchone()
    if not teacher_id_result:
        return 'Không tìm thấy thông tin về giáo viên.'

    teacher_id = teacher_id_result[0]
    teacher_name = teacher_id_result[1]

    # # Lấy các ClassID mà giáo viên đó kiểm soát từ bảng Classes
    # sql_teacher_classes = "SELECT ClassID FROM Classes WHERE TeacherID = %s"
    # mycursor.execute(sql_teacher_classes, (teacher_id,))
    # teacher_class_ids = [result[0] for result in mycursor.fetchall()]

    # Lấy các Email và Fullname tương ứng từ bảng Learners
    sql_class_learners = "SELECT LearnerID, Email, Fullname FROM Learners WHERE ClassID = %s"
    mycursor.execute(sql_class_learners, (class_id,))
    learner_info = mycursor.fetchall()


    # Tách riêng các giá trị Email và Fullname
    learner_ids = [info[0] for info in learner_info]
    emails = [info[1] for info in learner_info]
    fullnames = [info[2] for info in learner_info]

    fromdate = session.get('fromdate', '')
    todate = session.get('todate', '')

    learner_type_1_counts = {}
    for learner_id in learner_ids:
        sql_count_type_1 = """
            SELECT COUNT(*) AS Type1Count
            FROM code_results
            WHERE LearnerID = %s AND Type = 1 AND Date BETWEEN %s AND %s
        """
        mycursor.execute(sql_count_type_1, (learner_id,fromdate,todate))
        result = mycursor.fetchone()
        learner_type_1_counts[learner_id] = result[0]

    learner_type_0_counts = {}
    for learner_id in learner_ids:
        sql_count_type_0 = """
            SELECT COUNT(*) AS Type0Count
            FROM code_results
            WHERE LearnerID = %s AND Type = 0 AND Date BETWEEN %s AND %s
        """
        mycursor.execute(sql_count_type_0, (learner_id,fromdate,todate))
        result = mycursor.fetchone()
        learner_type_0_counts[learner_id] = result[0]

    learner_average_times = {}
    for learner_id in learner_ids:
        response = requests.get(url_for('average_completion_time', learner_id=learner_id, _external=True))
        data = response.json()
        learner_average_times[data['learner_id']] = data['average_completion_time']

    mycursor.execute("SELECT * FROM screenshots WHERE LearnerId = %s", ("2",))
    filenamesimg = mycursor.fetchall()

    return render_template('control2.html', results=results, result=result, emails=emails, fullnames=fullnames, learner_ids=learner_ids, learner_type_1_counts=learner_type_1_counts, teacher_name=teacher_name, learner_type_0_counts= learner_type_0_counts, learner_average_times=learner_average_times, filenamesimg= filenamesimg, class_id= class_id)


def get_result2():
    code = session.get('code', '')
    
    # Sử dụng API của ChatGPT để phân tích đề của code
    prompt = f"Đưa ra đề bài cho code sau đây(1 dòng)(bằng tiếng việt):\n\n{code}\n\n"

    response = openai.chat.completions.create(
        model="gpt-3.5-turbo", 
        messages=[
            {"role": "system", "content": prompt},
        ],
        max_tokens=1024
    )

    problem_statement = response.choices[0].message.content.strip()

   # Lưu kết quả, thời gian, và số lần click vào danh sách
    result = session.get('result', '')

    results.append({'code': code, 'problem_statement': problem_statement, 'result':result})

    # Gửi kết quả đến tất cả các clients đang kết nối
    socketio.emit('update_result', {'code': code, 'problem_statement': problem_statement, 'result': result})

    session['problem_statement'] = problem_statement

##------------------------------------------------------------------------------------------------------------##
    # Lấy thông tin về email của giáo viên đang đăng nhập từ session


    return render_template('control2.html', code=code, problem_statement=problem_statement, results=results, result=result)

def check_for_file_jdoodle(result):
    if "File \"jdoodle.py\"" in result or "main.c" in result:
        return 1
    else:
        return 0
    
@app.route("/it_mode", methods=['GET', 'POST'])
def it_mode():
    learner_id_2 = session.get('acc', '')
    mycursor.execute("SELECT * FROM screenshots WHERE LearnerId = %s", ("2",))
    filenamesimg = mycursor.fetchall()
    sql_get_code = "SELECT ID, Message FROM code_results WHERE LearnerId = %s AND Type = 0"
    mycursor.execute(sql_get_code, (learner_id_2,))
    files = mycursor.fetchall()
    username = session.get('fullname', '')
    if request.method == 'POST':
        code = request.form['code']
        session['code'] = code

        result = request.form['result']
        session['result'] = result

        problem_statement = session.get('problem_statement', '')


        socketio.emit('update_result', {'code': code,  'result': result})
        result_type = check_for_file_jdoodle(result)

        now = datetime.now()
        socketio.emit('update_datetime', {'date': now.strftime('%d/%m/%Y'), 'time': now.strftime('%H:%M:%S')})
        
        prompt = f"Phân tích ra đề bài của code sau (chỉ 1 dòng) (bằng tiếng việt):\n\n{code}\n\n"

        response = openai.chat.completions.create(
            model="gpt-3.5-turbo", 
            messages=[
                {"role": "system", "content": prompt},
            ],
            max_tokens=1024
        )   

        problem_statement = response.choices[0].message.content.strip()

    # Gửi kết quả đến tất cả các clients đang kết nối
        socketio.emit('update_result', {'code': code, 'problem_statement': problem_statement, 'result': result})

        session['problem_statement'] = problem_statement

        if "File \"jdoodle.py\"" in result or "main.c" in result:
            save_to_database(learner_id_2,result,code, now, result_type)
        else:
            save_to_database(learner_id_2,problem_statement,code, now, result_type)
     
    return render_template("it-mode.html",files=files, username= username, filenamesimg= filenamesimg)

@app.route('/open/<int:file_id>')
def open_file(file_id):
    try:
        # Truy vấn để lấy dữ liệu code của file từ cơ sở dữ liệu
        sql_get_code = "SELECT Code FROM code_results WHERE Id = %s"
        mycursor.execute(sql_get_code, (file_id,))
        file_data = mycursor.fetchone()

        # Kiểm tra xem có dữ liệu trả về không
        if file_data:
            code = file_data[0]
            return jsonify({'code': code})
        else:
            return jsonify({'error': 'File not found'})
    except Exception as e:
        return jsonify({'error': str(e)})
    
@app.route('/chart/<learner_id>/<from_date>/<to_date>')
def chart(learner_id, from_date, to_date):
    try:
        # Truy vấn để lấy thông tin về thời gian hoàn thành các bài học
        sql_get_completion_time = """
            SELECT ID, Type, Time
            FROM code_results
            WHERE LearnerID = %s AND Date BETWEEN %s AND %s
            ORDER BY ID;
        """
        mycursor.execute(sql_get_completion_time, (learner_id, from_date, to_date))
        results = mycursor.fetchall()

        # Xử lý dữ liệu để tính thời gian hoàn thành của mỗi bài học
        completion_times = []
        start_time = None
        count = 1
        for result in results:
            if result[1] == 1:  # Nếu là type 1 (lỗi)
                start_time = result[2]
            elif result[1] == 0 and start_time is not None:  # Nếu là type 0 (hoàn thành)
                end_time = result[2]
                time_difference = end_time - start_time
                completion_times.append(("Bài " + str(count), time_difference.total_seconds() // 60))  # Chuyển sang phút và lưu lại
                count += 1
                start_time = None

        # Tách ra danh sách ID bài học và thời gian hoàn thành
        ids = [data[0] for data in completion_times]
        completion_minutes = [data[1] for data in completion_times]

        return jsonify({'ids': ids, 'completion_minutes': completion_minutes})
    except Exception as e:
        return str(e)

@app.route('/date_range/<learner_id>')
def get_date_range(learner_id):
    try:
        # Lấy fromDate của ID 1
        sql_get_from_date = """
            SELECT MIN(Date)
            FROM code_results
            WHERE LearnerID = %s;
        """
        mycursor.execute(sql_get_from_date, (learner_id,))
        from_date = mycursor.fetchone()[0]

        # Lấy toDate của ID lớn nhất
        sql_get_to_date = """
            SELECT MAX(Date)
            FROM code_results
            WHERE LearnerID = %s;
        """
        mycursor.execute(sql_get_to_date, (learner_id,))
        to_date = mycursor.fetchone()[0]

        return jsonify({'from_date': from_date, 'to_date': to_date})
    except Exception as e:
        return str(e)

@app.route('/bar_chart/<learner_id>/<from_date>/<to_date>')
def bar_chart(learner_id, from_date, to_date):
    try:
        # Truy vấn để lấy thông tin về số lượng lỗi của từng bài
        sql_get_error_count = """
            SELECT ID, Type
            FROM code_results
            WHERE LearnerID = %s AND Date BETWEEN %s AND %s
            ORDER BY ID;
        """
        mycursor.execute(sql_get_error_count, (learner_id,from_date, to_date))
        results2 = mycursor.fetchall()

        # Xử lý dữ liệu để đếm số lượng lỗi của từng bài
        error_counts = []
        error_count = 0
        for result in results2:
            if result[1] == 1:  # Nếu là type 1 (lỗi)
                error_count += 1
            elif result[1] == 0 and error_count > 0:  # Nếu là type 0 (chạy đúng) và đã có lỗi được đếm
                error_counts.append(error_count)  # Lưu số lượng lỗi vào danh sách
                error_count = 0  # Reset lại biến đếm

        # Tạo danh sách số thứ tự bài học
        idss = ["Bài " + str(i + 1) for i in range(len(error_counts))]

        return jsonify({'idss': idss, 'error_counts': error_counts})
    except Exception as e:
        return str(e)

@app.route('/pie_chart_data/<learner_id>')
def pie_chart_data(learner_id):
    try:
        # Truy vấn để đếm số hàng có Type bằng 0 và bằng 1
        sql_count_type_1 = """
            SELECT 
                SUM(CASE WHEN Type = 0 THEN 1 ELSE 0 END) AS Type0Count,
                SUM(CASE WHEN Type = 1 THEN 1 ELSE 0 END) AS Type1Count
            FROM code_results
            WHERE LearnerID = %s;
        """
        mycursor.execute(sql_count_type_1, (learner_id,))
        type_counts = mycursor.fetchone()

        # Trả về dữ liệu dưới dạng JSON
        return jsonify({'type0Count': type_counts[0], 'type1Count': type_counts[1]})
    except Exception as e:
        return str(e)

@app.route('/get_advice2', methods=['POST'])
def get_advice2():
    data = request.json
    fullname = data['fullname']
    type1Count = data['type1Count']
    type0Count = data['type0Count']
    averageTime = data['averageTime']

    prompt =  f"Đưa ra lời nhận xét học tập lập trình với học viên '{fullname}' với số bài làm được là '{type0Count}' cùng số lỗi nhận được là '{type1Count}', cùng với tốc độ trung bình để hoàn thành là '{averageTime}', lời nhận xét có độ dài khá dài"

    # Use the OpenAI ChatGPT API to get a response
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo", 
        messages=[
			{"role": "system", "content": prompt},
		],
        max_tokens=1024
    )

    advice = response.choices[0].message.content.strip()

    return jsonify({'advice': advice})

@app.route('/get_advice3', methods=['POST'])
def get_advice3():
    data = request.json
    fullname = data['fullname']

    prompt =  f"tôi có một file data.txt để ghi lại những sự kiện được gõ trên bàn phím người dùng, Yêu cầu: phân tích nội dụng file dưới đây và cho biết người dùng đã làm những việc gì trong hoạt động file ghi lại ( có thể đưa ra một số dự đoán kết hợp dự đoán trang web người dùng sẽ vào. Nếu thấy địa chỉ hay tên trang của một trò chơi hay một trang web không liên quan trực tiếp tới việc học tập như facebook youtube,... thì cảnh báo học sinh). Nội dung file như sau: '{fullname}', (No Yapping)"

    # Use the OpenAI ChatGPT API to get a response
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo", 
        messages=[
			{"role": "system", "content": prompt},
		],
        max_tokens=1024
    )

    advice = response.choices[0].message.content.strip()
    session['advice'] = advice

    return jsonify({'advice': advice})

@app.route('/piechart')
def piechart():
    try:
        input_text = session.get('advice', '')
        percentages_int = re.findall(r'(\d+(?:\.\d+)?)%', input_text)
        percentages = percentages_int[:2]  # Only take the first two percentages
        percentages = [float(percentage.replace('%', '')) for percentage in percentages]
        return jsonify({'percentages': percentages})
    except Exception as e:
        return str(e)

@app.route("/main", methods=['GET', 'POST'])
def main():
    learner_id_2 = session.get('acc', '')
    if request.method == 'POST':
        code = request.form['code']
        session['code'] = code

        result = request.form['result']
        session['result'] = result

        problem_statement = session.get('problem_statement', '')

        socketio.emit('update_result', {'code': code,  'result': result})
        result_type = check_for_file_jdoodle(result)

        now = datetime.now()
        socketio.emit('update_datetime', {'date': now.strftime('%d/%m/%Y'), 'time': now.strftime('%H:%M:%S')})
        if "File \"jdoodle.py\"" in result or "main.c" in result:
            save_to_database(learner_id_2,result, now, result_type)
        else:
            save_to_database(learner_id_2,problem_statement, now, result_type)
        
        return get_result2()
    return render_template("main.html")

@app.route("/setting")
@login_required
def setting():
    return render_template("setting.html")

@app.route("/key_word_page")
@login_required
def key_word_page():
    return render_template("key_word_page.html")

@app.route("/control2/<class_id>", methods=['GET', 'POST'])
@login_required
def control2(class_id):
    fromdate="2024-02-20"
    todate="2024-03-20"
    if request.method == 'POST':
        fromdate = request.form['fromdate']
        todate = request.form['todate']
       
    session['fromdate'] = fromdate
    session['todate'] = todate
    session['classid'] = class_id
    return get_result(class_id)
####### DATABSE ############
def save_to_database(learner_id, error, code, now, type_value):
    sql = "INSERT INTO code_results (LearnerID, Type, Message, Code, Date, Time) VALUES (%s, %s, %s, %s, %s, %s)"
    val = (learner_id, type_value, error, code, now.strftime('%Y-%m-%d'), now.strftime('%H:%M:%S'))
    mycursor.execute(sql, val)
    mydb.commit()
    print(mycursor.rowcount, "record inserted.")

####### DATABSE ############
@app.route("/word")
def word():
    return render_template("word.html")


timer_duration = 10  # Set the timer duration (in seconds)

@app.route("/send_notification_with_image", methods=["POST"])
def send_notification_with_image():
    try:
         # Đọc nội dung tệp HTML từ tệp email_content.html
        email_content_file = os.path.join(app.root_path, "templates/email_content.html")
        with open(email_content_file, "r", encoding="utf-8") as file:
            html_content = file.read()
        
        # Gửi thông báo qua email
        from_email = "fluxiolmorax@gmail.com"
        to_email = "fluxiolmorax@gmail.com"
        app_password = "ydgs ohmo pklq ppfv"  # Sử dụng Mật khẩu Ứng dụng của bạn
        
        message = MIMEMultipart()
        message["From"] = from_email
        message["To"] = to_email
        message["Subject"] = "Thông báo mới"

        html = MIMEText(html_content, "html")
        message.attach(html)

        server = smtplib.SMTP("smtp.gmail.com", 587)
        server.starttls()
        server.login(from_email, app_password)

        server.sendmail(from_email, to_email, message.as_string())
        server.quit()

        return jsonify({"message": "Email đã được gửi thành công."})
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route("/test")
def test():
    return render_template("test.html")

# Twilio setup
account_sid = 'ACa0f31501e024b519906f9c22f8948d11'
auth_token = '03c258a540eb0436dc9b2ca95d8c5797'
client = Client(account_sid, auth_token)

# Email setup
from_email = "fluxiolmorax@gmail.com"
to_email = "fluxiolmorax@gmail.com"
app_password = "ydgs ohmo pklq ppfv"  # Use your App Password
message_sent = False

message_sent = False

@app.route('/focus')
def focus():
    global message_sent
    try:
        if not message_sent:
            # Read HTML content from the file
            email_content_file = os.path.join(app.root_path, "templates/email_tb.html")
            with open(email_content_file, "r", encoding="utf-8") as file:
                html_content = file.read()

            # Send email
            message = MIMEMultipart()
            message["From"] = from_email
            message["To"] = to_email
            message["Subject"] = "Cảnh Báo không chú ý học tập"

            html = MIMEText(html_content, "html")
            message.attach(html)

            server = smtplib.SMTP("smtp.gmail.com", 587)
            server.starttls()
            server.login(from_email, app_password)

            server.sendmail(from_email, to_email, message.as_string())
            server.quit()

            # Send Twilio SMS
            to_phone_number = '+84866139890'  # Replace with your actual phone number
            from_phone_number = '+12512902672'  # Replace with your Twilio phone number

            twilio_message = client.messages.create(
                body="Cảnh Báo không chú ý học tập",
                from_=from_phone_number,
                to=to_phone_number
            )

            message_sent = True

        return jsonify({"message": "Email and SMS sent successfully."})
    except Exception as e:
        return jsonify({"error": str(e)})
@app.route('/reset')
def reset():
    global message_sent
    message_sent = False
    return jsonify({"message": "Message reset successfully."})

UPLOAD_FOLDER = 'static/uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# def convert_doc_to_md(doc_file):
#     document = Document(doc_file)
#     md_content = ""
#     for paragraph in document.paragraphs:
#         md_content += paragraph.text + '\n'
#     return md_content
def convert_doc_to_md(doc_file):
    document = Document(doc_file)
    md_content = ""

    for paragraph in document.paragraphs:
        md_content += '<span style="'
        
        # Check if the paragraph is bold
        if any(run.bold for run in paragraph.runs):
            md_content += 'color: #FFF; font-family: \'Fira Code\'; font-size: 13px; font-style: normal; font-weight: 700; line-height: normal;">'
        else:
            md_content += 'color: rgba(255, 255, 255, 0.78); font-family: Fira Code; font-size: 12px; font-style: normal; font-weight: 500; line-height: 17px;">'
        for run in paragraph.runs:
            md_content += run.text

        md_content += '</span><br>'

    return md_content

def convert_markdown(text):
    extensions = ['extra', 'smarty', 'fenced_code', 'codehilite', 'tables', 'attr_list', 'def_list', 'abbr', 'footnotes', 'admonition', 'toc', 'meta', 'nl2br', 'sane_lists', 'toc', 'wikilinks']
    md = markdown.Markdown(extensions=extensions)
    html_content = md.convert(text)
    return html_content

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return "No file part"

    file = request.files['file']

    if file.filename == '':
        return "No selected file"

    file_path = f"{app.config['UPLOAD_FOLDER']}/{file.filename}"
    file.save(file_path)
    if file.filename.lower().endswith(('.doc', '.docx')):
        # Convert DOC/DOCX to Markdown
        md_content = convert_doc_to_md(file_path)
        # Convert Markdown if needed
        html_content = convert_markdown(md_content)

        return render_template('main.html', html_content=html_content)
    
@app.route('/upload2', methods=['POST'])
def upload2():
    if 'file' not in request.files:
        return "No file part"

    file = request.files['file']

    if file.filename == '':
        return "No selected file"

    file_path = f"{app.config['UPLOAD_FOLDER']}/{file.filename}"
    file.save(file_path)
    if file.filename.lower().endswith(('.doc', '.docx')):
        # Convert DOC/DOCX to Markdown
        md_content = convert_doc_to_md(file_path)
        # Convert Markdown if needed
        html_content = convert_markdown(md_content)

        return render_template('setting4_result.html', html_content=html_content)
    # if file.filename.lower().endswith(('.jpg', '.jpeg', '.png', '.gif')):
    #     # Display the image directly
    #     return render_template('setting2-result.html', image_path=file_path)
    # else:
    #     return "Vui Lòng Chọn File"

# Set your OpenAI GPT API key here
# Chat history
def get_completion(prompt): 
    print(prompt) 
    query = openai.chat.completions.create(
        model="gpt-3.5-turbo", 
        messages=[
			{"role": "system", "content": prompt},
		],
        max_tokens=1024,
        n=1,
        stop=None,
        temperature=0.5
    )

    response = query.choices[0].message.content.strip()
    return response 
  
@app.route("/get_response", methods=['POST', 'GET']) 
def query_view(): 
    if request.method == 'POST': 
        print('step1') 
        prompt = request.form['prompt'] 
        response = get_completion(prompt) 
        print(response) 
  
        return jsonify({'response': response}) 
    return render_template('test.html') 

jdoodle_api_url = 'https://api.jdoodle.com/v1/execute'

# JDoodle API credentials
client_id = '43c5e090725047f84652f38e3151160f'
client_secret = '6f6f43c340c0f3350f5b45a294cb5ac21aee63cc5ac40c4a0bac39256e93632a'
@app.route('/compile', methods=['POST'])
def compile_code():
    try:
        code = request.form['code']
        language = request.form['language']
        stdin = request.form['stdin']

        # JDoodle API request headers
        headers = {
            'Content-Type': 'application/json',
        }

        # JDoodle API request payload
        payload = {
            'clientId': client_id,
            'clientSecret': client_secret,
            'script': code,
            'language': language,
            'versionIndex': 0,  # Use '0' for the latest version of the language
            'stdin': stdin
        }

        # Send a request to the JDoodle Compiler API
        response = requests.post(jdoodle_api_url, json=payload, headers=headers)
        result = response.json()

        result['result'] = result['output']

        return jsonify(result)

    except Exception as e:
        return jsonify({'error': str(e)})

# -------------------------------------------------------Vô lý ----------------------------------
app.secret_key = b'_5#y2L"F4Q8z\n\xec]/'
run_threads = False

def block_keys():
    blocked_keys = {'alt+tab', 'f11'}

    while run_threads:
        try:
            event = keyboard.read_event(suppress=True)

            if f"{event.name.lower()}+{event.event_type}" in blocked_keys:
                print(f"Blocking key {event.name}")

                if f"{event.name.lower()}+{event.event_type}" == 'alt+tab':
                    keyboard.press_and_release('alt+esc')

                continue

        except Exception as e:
            print("Error:", e)

@app.route('/run_code', methods=['POST'])
def run_code():
    global run_threads

    timeout_seconds = int(request.form.get('timeout', 60))

    run_threads = not run_threads

    if run_threads:
        mouse_thread = threading.Thread(target=check_and_handle_mouse_position)
        key_block_thread = threading.Thread(target=block_keys)

        mouse_thread.start()
        key_block_thread.start()

        session['timeout'] = timeout_seconds

        print("Button clicked!")

    return render_template('main.html', timeout=timeout_seconds)

@app.route('/stop_threads')
def stop_threads():
    global run_threads
    run_threads = False
    return jsonify({'message': 'Threads stopped successfully'})

# -------------------------------------------------------Vô lý ----------------------------------

@app.route('/submit', methods=['POST'])
def submit():
    emailphuhuynh = request.form.get('emailphuhuynh')

    if 'quản lý' in emailphuhuynh:
        return redirect('/control')
    elif 'học sinh' in emailphuhuynh:
        return redirect('/setting')
    else:
        return redirect('/default')



# Classes Table
mycursor.execute("SHOW TABLES LIKE 'Screenshots'")
result = mycursor.fetchone()
if not result:
    mycursor.execute("""
        CREATE TABLE Screenshots (
            ImgID INT AUTO_INCREMENT PRIMARY KEY,
            LearnerID INT,
            image_path VARCHAR(255),
            capture_time TIME,
            RoomID INT
        )
    """)
    print("Table 'Screen' created successfully.")
else:
    print("Table 'Screen' already exists.")

UPLOAD_FOLDER2 = 'screenshots'
app.config['UPLOAD_FOLDER2'] = UPLOAD_FOLDER2

# def create_screenshots_directory():
#     if not os.path.exists(app.config['UPLOAD_FOLDER2']):
#         os.makedirs(app.config['UPLOAD_FOLDER2'])

# def capture_screenshot():
#     screenshot = pyautogui.screenshot()
#     screenshot = screenshot.resize((960, 540)) 
#     timestamp = datetime.now().strftime("%Y-%m-%d %H-%M-%S")
#     filename = f'screenshot_{timestamp}.png'
#     screenshot_path = os.path.join(app.config['UPLOAD_FOLDER2'], filename)
#     screenshot.save(screenshot_path)
    # return screenshot_path
    # cursor.execute("INSERT INTO screenshots (learnerid, image_path, capture_time) VALUES (%s, %s, %s)", (learner_id_2, screenshot_path, timestamp))
    # db.commit()

# def get_screenshot_filenames():
#     return os.listdir(app.config['UPLOAD_FOLDER2'])

@app.route("/test2")
def test2():
    return render_template('test2.html')
@app.route("/control")
@login_required
def control():
    learner_id_2 = session.get('acc', '')
    cursor.execute("SELECT * FROM screenshots WHERE LearnerId = %s", (learner_id_2,))
    filenames = cursor.fetchall()
    return render_template("control.html", filenames=filenames)


@app.route('/screenshots/<filename>')
def get_screenshot(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER2'], filename)

@app.route('/data_get/<filename>')
def get_screenshot2(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER3'], filename)

# def auto_capture():
#     while True:
#         time.sleep(30)
#         capture_screenshot()
#         get_image(URL, SCREENSHOT)
CORS(app)


########### DATABASE ##############
db = mysql.connector.connect(
    host=os.environ.get('DATABASE_HOST'),
    user=os.environ.get('DATABASE_USER'),
    password=os.environ.get('DATABASE_PASS'),
    database=os.environ.get('DATABASE'),
)
cursor = db.cursor()

cursor.execute("""
    CREATE TABLE IF NOT EXISTS message (
        id INT AUTO_INCREMENT PRIMARY KEY,
        content VARCHAR(200) NOT NULL,
        ClassID INT
    )
""")

@app.route('/get_messages2', methods=['GET'])
def get_messages2():
    class_id = session.get('classid', '')
    cursor.execute("SELECT content FROM message WHERE ClassID = %s", (class_id,))
    messages = cursor.fetchall()
    messages_list = [{'content': message[0]} for message in messages]
    return jsonify({'messages': messages_list})

@app.route('/send_message2', methods=['POST'])
def send_message2():
    class_id = session.get('classid', '')
    content = request.json['content']
    cursor.execute("INSERT INTO message (content,classID) VALUES (%s, %s)", (content, class_id))
    db.commit()
    return jsonify({'status': 'Message sent successfully'})


messagess = []

@app.route('/send_message', methods=['POST'])
def send_message():
    global messagess
    message = request.json.get('message', '')
    messagess.append(message)
    return {'status': 'success'}

@app.route('/get_messages')
def get_messages():
    def generate():
        while True:
            if messagess:
                yield f"data: {messagess[-1]}\n\n"
                messagess.pop()
    return Response(generate(), content_type='text/event-stream')

# BASE_DIR = Path(__file__).resolve().parent

# URL = 'http://localhost:5000/api/upload_image'

# SCREENSHOT = BASE_DIR / 'screenshots'

# def get_image(url, directory):
#     screenshot_files = os.listdir(directory)
#     for filename in screenshot_files:
#         with open(directory / filename, 'rb') as img:
#             response = requests.post(url, files={'image': img})

@app.route("/join", methods=["GET", "POST"])
def join():
    # capture_thread = Thread(target=auto_capture)
    # capture_thread.start()
    learner_id_2 = session.get('acc', '')
    if request.method == "POST":
        room_id = request.form.get("roomID")
        cursor.execute("INSERT INTO screenshots (learnerID, roomID) VALUES (%s, %s)", (learner_id_2, room_id,))
        db.commit()
        return redirect(f"/it_mode?roomID={room_id}")
    return render_template("join.html")

@app.route("/meeting2_test")
@login_required
def meeting2():
    return render_template("meeting2_test.html")

@app.route("/join2", methods=["GET", "POST"])
@login_required
def join2():
    if request.method == "POST":
        room_id = request.form.get("roomID")
        return redirect(f"/meeting2_test?roomID={room_id}")
    return render_template("join2.html")

@app.route('/api/upload_image', methods=['POST'])
def get_image():
    img = request.files.get('image')
    filename=img.filename
    screenshot_path = os.path.join(app.config['UPLOAD_FOLDER2'], filename)
    img.save(screenshot_path)
    timestamp = datetime.now().strftime("%Y-%m-%d %H-%M-%S")
    cursor.execute("INSERT INTO screenshots (learnerID, image_path, capture_time) VALUES (%s, %s, %s)", ("2", screenshot_path, timestamp))
    db.commit()
    return "Success"

@app.route('/api/upload_data_key', methods=['POST'])
def data_key():
    data_key = request.files.get('file')
    filename=data_key.filename
    data_path = os.path.join(app.config['UPLOAD_FOLDER3'], filename)
    data_key.save(data_path)
    return "Success"


if __name__ == '__main__':
    app.run(debug=True)
    
